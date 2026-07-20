import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from config.styles import (
    PAGE_MARGIN, COL_WIDTHS, COLOR_PRIMARY_HEX, COLOR_GROUP_BG_HEX,
    COLOR_MAIN_TEXT_RGB, COLOR_SKIP_QUESTION_RGB, COLOR_HINT_RGB,
    COLOR_HEADER_TEXT_RGB, COLOR_GROUP_TEXT_RGB,
    FONT_SIZE_MAIN, FONT_SIZE_SKIP, FONT_SIZE_HINT, FONT_NAME
)

def set_cell_background(cell, fill_hex):
    """Applies shading to table cells."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def apply_font_settings(run, font_name=FONT_NAME, size=FONT_SIZE_MAIN, bold=False, italic=False, color=COLOR_MAIN_TEXT_RGB):
    """Applies consistent font configuration to a text run."""
    run.font.name = font_name
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color

def create_document(lang_name, label_col, survey_df, choices_dict, title, output_path):
    doc = Document()
    
    # Configure Margins
    section = doc.sections[0]
    section.top_margin = PAGE_MARGIN
    section.bottom_margin = PAGE_MARGIN
    section.left_margin = PAGE_MARGIN
    section.right_margin = PAGE_MARGIN

    # Document Header
    h = doc.add_heading(level=1)
    h_run = h.add_run(str(title))
    apply_font_settings(h_run, size=Pt(16), bold=True, color=RGBColor(30, 30, 30))
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sub = doc.add_paragraph()
    sub_run = sub.add_run(f"Language: {lang_name}")
    apply_font_settings(sub_run, size=Pt(10), italic=True, color=RGBColor(100, 100, 100))
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Table Setup
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Table Header Row
    hdr_cells = table.rows[0].cells
    hdr_titles = ["S.N.", "Question & Details", "Options / Input Format"]
    
    for idx, cell in enumerate(hdr_cells):
        set_cell_background(cell, COLOR_PRIMARY_HEX)
        cell.text = hdr_titles[idx]
        for p in cell.paragraphs:
            for run in p.runs:
                apply_font_settings(run, size=Pt(11), bold=True, color=COLOR_HEADER_TEXT_RGB)

    hint_col = f"hint::{lang_name}" if f"hint::{lang_name}" in survey_df.columns else None
    sn_counter = 0

    for _, row in survey_df.iterrows():
        q_type = str(row.get('type', '')).strip()

        # Skip technical / non-display fields
        if q_type in ['calculate', 'start', 'end', 'today', 'deviceid', 'phonenumber'] or pd.isna(q_type):
            continue

        # Section / Group Headers
        if q_type.startswith(('begin_group', 'begin_repeat')):
            group_label = row.get(label_col)
            if pd.isna(group_label) or not str(group_label).strip():
                group_label = row.get('name', 'Group')
                
            row_cells = table.add_row().cells
            row_cells[0].text = ""
            row_cells[1].text = f"SECTION: {group_label}".upper()
            row_cells[1].merge(row_cells[2])
            
            set_cell_background(row_cells[0], COLOR_GROUP_BG_HEX)
            set_cell_background(row_cells[1], COLOR_GROUP_BG_HEX)
            
            for p in row_cells[1].paragraphs:
                for run in p.runs:
                    apply_font_settings(run, size=Pt(11), bold=True, color=COLOR_GROUP_TEXT_RGB)
            continue

        if q_type.startswith(('end_group', 'end_repeat')):
            continue

        # Standard Questions
        sn_counter += 1
        q_label = row.get(label_col)
        if pd.isna(q_label) or not str(q_label).strip():
            q_label = str(row.get('name', f'Question_{sn_counter}'))

        q_hint = row.get(hint_col) if hint_col else None
        
        # Determine if this question depends on skip logic (has a 'relevant' field)
        has_skip_logic = pd.notna(row.get('relevant')) and str(row.get('relevant')).strip() != ""
        
        # Apply font size & color hierarchy
        current_font_size = FONT_SIZE_SKIP if has_skip_logic else FONT_SIZE_MAIN
        current_text_color = COLOR_SKIP_QUESTION_RGB if has_skip_logic else COLOR_MAIN_TEXT_RGB
        
        req_flag = " *" if str(row.get('required')).lower() in ['true', '1', 'yes'] else ""

        row_cells = table.add_row().cells
        
        # S.N. Column
        row_cells[0].text = str(sn_counter)
        p0 = row_cells[0].paragraphs[0]
        if len(p0.runs) > 0:
            apply_font_settings(p0.runs[0], size=current_font_size, color=current_text_color)

        # Question Column
        p2 = row_cells[1].paragraphs[0]
        
        # Apply indentation if the question is skip-logic dependent
        if has_skip_logic:
            p2.paragraph_format.left_indent = Inches(0.25)

        r2 = p2.add_run(f"{q_label}{req_flag}")
        apply_font_settings(r2, size=current_font_size, bold=True, color=current_text_color)

        if pd.notna(q_hint) and str(q_hint).strip():
            p2.add_run("\n")
            r_hint = p2.add_run(f"Hint: {str(q_hint).strip()}")
            apply_font_settings(r_hint, size=FONT_SIZE_HINT, italic=True, color=COLOR_HINT_RGB)

        # Options / Format Column
        p3 = row_cells[2].paragraphs[0]
        if has_skip_logic:
            p3.paragraph_format.left_indent = Inches(0.15)

        if q_type.startswith(('select_one', 'select_multiple')):
            parts = q_type.split()
            list_name = parts[1] if len(parts) > 1 else ""
            options = choices_dict.get(list_name, [])
            if options:
                p3.text = "\n".join([f"□  {opt}" for opt in options])
            else:
                p3.text = "□  [ Options List ]"
                
            for run in p3.runs:
                apply_font_settings(run, size=current_font_size, color=current_text_color)

        elif q_type == 'text':
            # Clean dotted response line
            r_text = p3.add_run("..................................................")
            apply_font_settings(r_text, size=current_font_size, color=current_text_color)

        elif q_type in ['integer', 'decimal']:
            r_num = p3.add_run("[ Number: ____________ ]")
            apply_font_settings(r_num, size=current_font_size, color=current_text_color)

        elif q_type == 'date':
            r_date = p3.add_run("[ Date: DD / MM / YYYY ]")
            apply_font_settings(r_date, size=current_font_size, color=current_text_color)

        elif q_type == 'time':
            r_time = p3.add_run("[ Time: HH : MM ]")
            apply_font_settings(r_time, size=current_font_size, color=current_text_color)

        elif q_type == 'note':
            r_note = p3.add_run("—")
            apply_font_settings(r_note, size=current_font_size, color=current_text_color)

        else:
            r_gen = p3.add_run(f"[ {q_type} ]")
            apply_font_settings(r_gen, size=current_font_size, color=current_text_color)

    # Set Explicit Table Widths
    for row in table.rows:
        for idx, width in enumerate(COL_WIDTHS):
            row.cells[idx].width = width

    doc.save(output_path)
    print(f"✔ Generated: {output_path}")